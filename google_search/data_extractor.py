"""
Data Extractor - Extract data from multiple file formats

Supports:
- HTML (BeautifulSoup)
- PDF (PyPDF2, pdfplumber)
- DOCX (python-docx)
- XLSX (openpyxl)
- JSON
- CSV
- TXT

Automatically detects format and extracts relevant data.
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
from abc import ABC, abstractmethod
import mimetypes


logger = logging.getLogger(__name__)


@dataclass
class ExtractedData:
    """Represents extracted data from a file."""
    source_file: Path
    file_type: str
    content: str
    metadata: Dict[str, Any]
    error: Optional[str] = None
    
    def is_successful(self) -> bool:
        """Check if extraction was successful."""
        return self.error is None
    
    def get_summary(self) -> str:
        """Get summary of extracted data."""
        return f"""
        File: {self.source_file.name}
        Type: {self.file_type}
        Content Length: {len(self.content)} chars
        Metadata: {self.metadata}
        Status: {'Success' if self.is_successful() else f'Error: {self.error}'}
        """


class FileExtractor(ABC):
    """Base class for file extractors."""
    
    @abstractmethod
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract data from file."""
        pass
    
    @abstractmethod
    def supports_format(self, file_type: str) -> bool:
        """Check if extractor supports this file type."""
        pass


class HTMLExtractor(FileExtractor):
    """Extract data from HTML files."""
    
    def supports_format(self, file_type: str) -> bool:
        return file_type.lower() in ['html', 'htm']
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html = f.read()
            
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()
            
            text = soup.get_text()
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            content = '\n'.join(lines)
            
            metadata = {
                'title': soup.title.string if soup.title else 'N/A',
                'language': soup.html.get('lang') if soup.html else 'N/A'
            }
            
            logger.info(f"Successfully extracted HTML from {file_path.name}")
            return ExtractedData(
                source_file=file_path,
                file_type='html',
                content=content,
                metadata=metadata
            )
            
        except ImportError:
            return ExtractedData(
                source_file=file_path,
                file_type='html',
                content='',
                metadata={},
                error='BeautifulSoup4 not installed: pip install beautifulsoup4'
            )
        except Exception as e:
            logger.error(f"Error extracting HTML: {e}")
            return ExtractedData(
                source_file=file_path,
                file_type='html',
                content='',
                metadata={},
                error=str(e)
            )


class PDFExtractor(FileExtractor):
    """Extract data from PDF files."""
    
    def supports_format(self, file_type: str) -> bool:
        return file_type.lower() == 'pdf'
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract text from PDF."""
        try:
            import pdfplumber
            
            content = []
            metadata = {'pages': 0, 'author': None}
            
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                if pdf.metadata:
                    metadata.update(pdf.metadata)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        content.append(f"--- Page {page_num} ---\n{text}")
            
            full_content = '\n'.join(content)
            
            logger.info(f"Successfully extracted PDF from {file_path.name}")
            return ExtractedData(
                source_file=file_path,
                file_type='pdf',
                content=full_content,
                metadata=metadata
            )
            
        except ImportError:
            return ExtractedData(
                source_file=file_path,
                file_type='pdf',
                content='',
                metadata={},
                error='pdfplumber not installed: pip install pdfplumber'
            )
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ExtractedData(
                source_file=file_path,
                file_type='pdf',
                content='',
                metadata={},
                error=str(e)
            )


class DOCXExtractor(FileExtractor):
    """Extract data from DOCX files."""
    
    def supports_format(self, file_type: str) -> bool:
        return file_type.lower() in ['docx', 'doc']
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            full_content = '\n'.join(content)
            
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            logger.info(f"Successfully extracted DOCX from {file_path.name}")
            return ExtractedData(
                source_file=file_path,
                file_type='docx',
                content=full_content,
                metadata=metadata
            )
            
        except ImportError:
            return ExtractedData(
                source_file=file_path,
                file_type='docx',
                content='',
                metadata={},
                error='python-docx not installed: pip install python-docx'
            )
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ExtractedData(
                source_file=file_path,
                file_type='docx',
                content='',
                metadata={},
                error=str(e)
            )


class XLSXExtractor(FileExtractor):
    """Extract data from XLSX files."""
    
    def supports_format(self, file_type: str) -> bool:
        return file_type.lower() in ['xlsx', 'xls']
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract data from XLSX."""
        try:
            from openpyxl import load_workbook
            
            workbook = load_workbook(file_path)
            content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else '' for cell in row]
                    content.append('\t'.join(row_data))
            
            full_content = '\n'.join(content)
            
            metadata = {
                'sheets': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames
            }
            
            logger.info(f"Successfully extracted XLSX from {file_path.name}")
            return ExtractedData(
                source_file=file_path,
                file_type='xlsx',
                content=full_content,
                metadata=metadata
            )
            
        except ImportError:
            return ExtractedData(
                source_file=file_path,
                file_type='xlsx',
                content='',
                metadata={},
                error='openpyxl not installed: pip install openpyxl'
            )
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
            return ExtractedData(
                source_file=file_path,
                file_type='xlsx',
                content='',
                metadata={},
                error=str(e)
            )


class JSONExtractor(FileExtractor):
    """Extract data from JSON files."""
    
    def supports_format(self, file_type: str) -> bool:
        return file_type.lower() == 'json'
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract data from JSON."""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            content = json.dumps(data, indent=2)
            
            metadata = {
                'keys': list(data.keys()) if isinstance(data, dict) else 'array',
                'size': len(content)
            }
            
            logger.info(f"Successfully extracted JSON from {file_path.name}")
            return ExtractedData(
                source_file=file_path,
                file_type='json',
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error extracting JSON: {e}")
            return ExtractedData(
                source_file=file_path,
                file_type='json',
                content='',
                metadata={},
                error=str(e)
            )


class CSVExtractor(FileExtractor):
    """Extract data from CSV files."""
    
    def supports_format(self, file_type: str) -> bool:
        return file_type.lower() == 'csv'
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract data from CSV."""
        try:
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                content = f.read()
            
            lines = content.split('\n')
            
            metadata = {
                'rows': len([l for l in lines if l.strip()]),
                'columns': len(lines[0].split(',')) if lines else 0
            }
            
            logger.info(f"Successfully extracted CSV from {file_path.name}")
            return ExtractedData(
                source_file=file_path,
                file_type='csv',
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error extracting CSV: {e}")
            return ExtractedData(
                source_file=file_path,
                file_type='csv',
                content='',
                metadata={},
                error=str(e)
            )


class TextExtractor(FileExtractor):
    """Extract data from TXT files."""
    
    def supports_format(self, file_type: str) -> bool:
        return file_type.lower() in ['txt', 'text']
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract data from TXT."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                'lines': len(content.split('\n')),
                'characters': len(content)
            }
            
            logger.info(f"Successfully extracted TXT from {file_path.name}")
            return ExtractedData(
                source_file=file_path,
                file_type='txt',
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            return ExtractedData(
                source_file=file_path,
                file_type='txt',
                content='',
                metadata={},
                error=str(e)
            )


class DataExtractor:
    """Main data extractor that uses appropriate extractor for each file type."""
    
    def __init__(self):
        self.extractors: List[FileExtractor] = [
            HTMLExtractor(),
            PDFExtractor(),
            DOCXExtractor(),
            XLSXExtractor(),
            JSONExtractor(),
            CSVExtractor(),
            TextExtractor()
        ]
        logger.info("DataExtractor initialized with 7 file format handlers")
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type from extension."""
        return file_path.suffix.lstrip('.').lower()
    
    def extract(self, file_path: Path) -> ExtractedData:
        """Extract data from file using appropriate extractor."""
        file_type = self._detect_file_type(file_path)
        
        for extractor in self.extractors:
            if extractor.supports_format(file_type):
                return extractor.extract(file_path)
        
        logger.warning(f"No extractor found for file type: {file_type}")
        return ExtractedData(
            source_file=file_path,
            file_type=file_type,
            content='',
            metadata={},
            error=f'Unsupported file format: {file_type}'
        )
    
    def extract_batch(self, directory: Path, pattern: str = '*') -> List[ExtractedData]:
        """Extract data from multiple files in a directory."""
        results = []
        
        for file_path in directory.glob(pattern):
            if file_path.is_file():
                result = self.extract(file_path)
                results.append(result)
        
        logger.info(f"Extracted data from {len(results)} files")
        return results
    
    def save_extracted_data(self, data: ExtractedData, output_dir: Path) -> Path:
        """Save extracted data to file."""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_file = output_dir / f"{data.source_file.stem}_extracted.txt"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"Source: {data.source_file.name}\n")
            f.write(f"Type: {data.file_type}\n")
            f.write(f"Metadata: {data.metadata}\n")
            f.write("="*80 + "\n\n")
            f.write(data.content)
        
        logger.info(f"Saved extracted data to {output_file}")
        return output_file
